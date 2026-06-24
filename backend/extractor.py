import io
import os
import re
import shutil
from collections import Counter
from typing import Callable, List, Optional, Tuple

try:
    import numpy as np  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    np = None

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    fitz = None

try:
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    pdfplumber = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None
    Image = None

try:
    from rapidocr_onnxruntime import RapidOCR  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    RapidOCR = None

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    Document = None


def normalize_extracted_text(text: str) -> str:
    text = text or ""
    text = (
        text.replace("\x00", " ")
        .replace("\u200b", "")
        .replace("\ufeff", "")
        .replace("\xa0", " ")
        .replace("\u2502", " | ")
        .replace("\u2010", "-")
        .replace("\u2011", "-")
        .replace("\u2012", "-")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2015", "-")
        .replace("\u2212", "-")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _dedupe_lines(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    seen = Counter()
    output = []
    for line in lines:
        if not line:
            continue
        normalized = re.sub(r"\s+", " ", line).strip().lower()
        # keep the same line max twice to avoid repeated table/ocr noise.
        if seen[normalized] >= 2:
            continue
        seen[normalized] += 1
        output.append(line)
    return "\n".join(output)


def _text_quality_score(text: str) -> float:
    clean = normalize_extracted_text(text)
    if not clean:
        return 0.0
    words = re.findall(r"\w+", clean)
    lines = [line.strip() for line in clean.splitlines() if line.strip()]
    if not words or not lines:
        return 0.0

    total_words = len(words)
    word_score = min(total_words / 260.0, 1.0) * 55.0

    single_char_lines = sum(
        1
        for line in lines
        if len(re.findall(r"\w+", line)) == 1 and len(re.sub(r"\W", "", line)) <= 2
    )
    short_fragment_ratio = single_char_lines / max(len(lines), 1)
    fragment_penalty = min(short_fragment_ratio * 45.0, 35.0)

    alpha_words = sum(1 for token in words if re.search(r"[A-Za-z]{2,}", token))
    alpha_ratio = alpha_words / max(total_words, 1)
    alpha_score = alpha_ratio * 20.0

    date_count = len(re.findall(r"\b(?:19|20)\d{2}\b", clean))
    date_score = min(date_count, 8) * 1.5

    control_chars = len(re.findall(r"[\u0001-\u0008\u000b\u000c\u000e-\u001f]", clean))
    control_penalty = 15.0 if control_chars else 0.0

    score = word_score + alpha_score + date_score - fragment_penalty - control_penalty
    return max(0.0, min(score, 100.0))


def _extract_with_pymupdf(file_path: str) -> str:
    if fitz is None:
        return ""
    chunks: List[str] = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                blocks = page.get_text("blocks") or []
                if blocks:
                    ordered = sorted(
                        blocks,
                        key=lambda item: (round(float(item[1]), 1), round(float(item[0]), 1)),
                    )
                    block_texts = [str(item[4] or "").strip() for item in ordered if str(item[4] or "").strip()]
                    if block_texts:
                        chunks.append("\n".join(block_texts))
                        continue
                page_text = page.get_text("text", sort=True) or ""
                if page_text.strip():
                    chunks.append(page_text)
    except Exception:
        return ""
    return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))


def _extract_with_pdfplumber(file_path: str) -> str:
    if pdfplumber is None:
        return ""
    chunks: List[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                if page_text.strip():
                    chunks.append(page_text)

                if len(re.findall(r"\w+", page_text)) >= 40:
                    continue

                try:
                    for table in page.extract_tables() or []:
                        for row in table or []:
                            cells = [str(cell or "").strip() for cell in row]
                            non_empty = [cell for cell in cells if cell]
                            if len(non_empty) < 2:
                                continue
                            joined = " | ".join(non_empty)
                            if joined and joined not in page_text:
                                chunks.append(joined)
                except Exception:
                    pass
    except Exception:
        return ""
    return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))


def _extract_with_pypdf(file_path: str) -> str:
    if PdfReader is None:
        return ""
    chunks: List[str] = []
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text)
    except Exception:
        return ""
    return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))


def _extract_with_ocr(file_path: str) -> str:
    if fitz is None or pytesseract is None or Image is None:
        return ""

    # Allow OCR to work even when tesseract is installed but not added to PATH.
    try:
        configured = os.getenv("TESSERACT_CMD", "").strip()
        if configured and os.path.exists(configured):
            pytesseract.pytesseract.tesseract_cmd = configured
        elif not shutil.which("tesseract"):
            default_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]
            for candidate in default_paths:
                if os.path.exists(candidate):
                    pytesseract.pytesseract.tesseract_cmd = candidate
                    break
    except Exception:
        pass

    chunks: List[str] = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                try:
                    # 2x zoom for better OCR quality.
                    matrix = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    text = pytesseract.image_to_string(image) or ""
                    if text.strip():
                        chunks.append(text)
                except Exception:
                    continue
    except Exception:
        return ""
    return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))


def _extract_with_rapidocr(file_path: str) -> str:
    if fitz is None or RapidOCR is None or np is None or Image is None:
        return ""
    chunks: List[str] = []
    try:
        engine = RapidOCR()
        with fitz.open(file_path) as doc:
            for page in doc:
                try:
                    matrix = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    image_bytes = pix.tobytes("png")
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB") if Image is not None else None
                    img_array = np.array(image) if image is not None else None
                    if img_array is None:
                        continue
                    result, _ = engine(img_array)
                    if not result:
                        continue
                    ordered = sorted(
                        result,
                        key=lambda row: (
                            float(row[0][0][1]) if row and row[0] and row[0][0] else 0.0,
                            float(row[0][0][0]) if row and row[0] and row[0][0] else 0.0,
                        ),
                    )
                    lines = [str(row[1] or "").strip() for row in ordered if len(row) > 1 and str(row[1] or "").strip()]
                    if lines:
                        chunks.append("\n".join(lines))
                except Exception:
                    continue
    except Exception:
        return ""
    return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))


def _select_best_extraction(candidates: List[Tuple[str, str]]) -> str:
    best_text = ""
    best_score = -1.0
    for _name, text in candidates:
        score = _text_quality_score(text)
        if score > best_score:
            best_score = score
            best_text = text

    if best_score < 30.0 and len(candidates) >= 2:
        # On very noisy docs combine top two candidates to preserve missing fields.
        scored = sorted(
            [(_text_quality_score(text), text) for _, text in candidates if text.strip()],
            key=lambda pair: pair[0],
            reverse=True,
        )
        if len(scored) >= 2:
            merged = _dedupe_lines(normalize_extracted_text(scored[0][1] + "\n" + scored[1][1]))
            if _text_quality_score(merged) >= best_score:
                return merged
    return best_text


def extract_text_from_pdf(file_path: str) -> str:
    extractors: List[Tuple[str, Callable[[str], str]]] = [
        ("pymupdf_layout", _extract_with_pymupdf),
        ("pdfplumber", _extract_with_pdfplumber),
        ("pypdf", _extract_with_pypdf),
    ]

    candidates: List[Tuple[str, str]] = []
    for name, extractor in extractors:
        text = extractor(file_path)
        if text.strip():
            candidates.append((name, text))

    # OCR fallback only when non-OCR extraction quality is poor.
    if not candidates or max(_text_quality_score(text) for _, text in candidates) < 55.0:
        rapid_ocr_text = _extract_with_rapidocr(file_path)
        if rapid_ocr_text.strip():
            candidates.append(("rapidocr", rapid_ocr_text))
        ocr_text = _extract_with_ocr(file_path)
        if ocr_text.strip():
            candidates.append(("ocr", ocr_text))

    if not candidates:
        return ""
    return _select_best_extraction(candidates)


def extract_text_from_docx(file_path: str) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(file_path)
        chunks = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return _dedupe_lines(normalize_extracted_text("\n".join(chunks)))
    except Exception:
        return ""


def extract_text_from_txt(file_path: str) -> str:
    encodings = ("utf-8-sig", "utf-8", "cp1252", "latin-1")
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as handle:
                return _dedupe_lines(normalize_extracted_text(handle.read()))
        except Exception:
            continue
    return ""


def extract_name(text: str) -> Optional[str]:
    if not text:
        return None

    lines = text.split("\n")
    for line in lines[:12]:
        line = line.strip()
        if len(line) < 3:
            continue
        line = re.split(r",|\||\(", line)[0]
        line = re.sub(r"\b(Dr\.?|Mr\.?|Mrs\.?|Ms\.?)\b", "", line, flags=re.IGNORECASE)
        line = re.sub(r"\s+", " ", line).strip()
        if len(line.split()) >= 2 and len(line) < 60:
            return line
    return None
